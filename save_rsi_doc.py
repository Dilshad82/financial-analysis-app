import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# إنشاء المستند
doc = docx.Document()

# العنوان الرئيسي
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("📊 دليل مؤشر القوة النسبية (RSI) - Relative Strength Index")
run.font.name = 'Arial'
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph("---------------------------------------------------------------------------------------------------------")

# المقدمة
p1 = doc.add_paragraph()
run1 = p1.add_run("ما هو مؤشر RSI؟\n")
run1.font.bold = True
run1.font.size = Pt(14)
p1.add_run("مؤشر القوة النسبية (RSI) هو أحد أشهر المؤشرات الفنية المستخدمة في أسواق المال (الذهب، الأسهم، العملات). يقيس سرعة وقوة التغير في أسعار الأصل المالي خلال فترة زمنية محددة (عادة 14 يوماً).\n\nتتراوح قيمة المؤشر بين 0 و 100، وتُحدد ما إذا كان السعر في منطقة تشبع شرائي أم تشبع بيعي.")

# مستويات المؤشر
p2 = doc.add_paragraph()
run2 = p2.add_run("\n📌 مستويات مؤشر RSI وكيفية قراءتها:\n")
run2.font.bold = True
run2.font.size = Pt(14)

doc.add_paragraph("1️⃣ منطقة التشبع الشرائي (Overbought) — فوق 70:\nتعني أن السعر ارتفع بشكل سريع ومبالغ فيه، ويُتوقع احتمال حدوث تصحيح هبوطي (فرصة بيع / جني أرباح).")
doc.add_paragraph("2️⃣ منطقة التشبع البيعي (Oversold) — تحت 30:\nتعني أن عمليات البيع كانت مكثفة للهبوط الحاد جداً، ويُتوقع قرب ارتداد صعودي (فرصة شراء محتملة).")
doc.add_paragraph("3️⃣ المنطقة المحايدة (Neutral Area) — بين 30 و 70:\nالسعر يتحرك بشكل طبيعي. يُعتبر المستوى 50 هو الخط الفاصل بين الاتجاه الصاعد والنازل.")

# التطبيق العملي
p3 = doc.add_paragraph()
run3 = p3.add_run("\n⚙️ كيف يعمل RSI داخل منصة التحليل؟\n")
run3.font.bold = True
run3.font.size = Pt(14)

p3.add_run("يتم دمج مؤشر RSI مع المتوسط المتحرك (SMA 20) لتأكيد التوصيات:\n")
doc.add_paragraph("• 🟢 إشارة شراء (BUY): عندما يكون السعر فوق المتوسط 20 و RSI أقل من 45.")
doc.add_paragraph("• 🔴 إشارة بيع (SELL): عندما يكون السعر تحت المتوسط 20 و RSI أعلى من 55.")

# حفظ الملف
file_name = "دليل_مؤشر_RSI.docx"
doc.save(file_name)
print(f"تم حفظ جميع السطور بنجاح في ملف: {file_name}")