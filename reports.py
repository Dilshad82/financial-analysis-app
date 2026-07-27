import io
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import pandas as pd


class ReportGenerator:

    @staticmethod
    def convert_df_to_excel(df: pd.DataFrame) -> bytes:
        """تصدير بيانات الأسعار والمؤشرات إلى ملف Excel"""
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Data')
        return output.getvalue()

    @staticmethod
    def generate_rsi_word_doc() -> bytes:
        """توليد ملف Word منسق يحتوي على شرح مؤشر RSI والمنظومة المالية الشاملة"""
        doc = docx.Document()

        # العنوان الرئيسي
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(
            "📊 دليل مؤشر القوة النسبية (RSI) والمنظومة المالية"
        )
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)

        doc.add_paragraph(
            "---------------------------------------------------------------------------------------------------------"
        )

        # المقدمة
        p1 = doc.add_paragraph()
        run1 = p1.add_run("ما هو مؤشر RSI؟\n")
        run1.font.name = 'Arial'
        run1.font.bold = True
        run1.font.size = Pt(14)
        run1.font.color.rgb = RGBColor(0, 51, 102)

        p1.add_run(
            "مؤشر القوة النسبية (Relative Strength Index - RSI) هو أحد أشهر المؤشرات الفنية المستخدمة في أسواق المال "
            "(الأسهم، العملات الرقمية، والذهب). يقيس سرعة وقوة التغير في أسعار الأصل المالي خلال فترة زمنية محددة (عادة 14 يوماً).\n\n"
            "تتراوح قيمة المؤشر بين 0 و 100، وتُحدد ما إذا كان السعر يمر بمرحلة تشبع شرائي أم تشبع بيعي."
        )

        # المستويات الفنية
        p2 = doc.add_paragraph()
        run2 = p2.add_run("\n📌 مستويات مؤشر RSI وكيفية قراءتها:\n")
        run2.font.name = 'Arial'
        run2.font.bold = True
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph(
            "1️⃣ منطقة التشبع الشرائي (Overbought) — فوق 70:\n"
            "تعني أن السعر ارتفع بشكل سريع ومبالغ فيه، ويُتوقع احتمال حدوث تصحيح هبوطي (فرصة بيع / جني أرباح)."
        )
        doc.add_paragraph(
            "2️⃣ منطقة التشبع البيعي (Oversold) — تحت 30:\n"
            "تعني أن عمليات البيع كانت مكثفة والهبوط حاد جداً، ويُتوقع قرب ارتداد صعودي (فرصة شراء محتملة)."
        )
        doc.add_paragraph(
            "3️⃣ المنطقة المحايدة (Neutral Area) — بين 30 و 70:\n"
            "السعر يتحرك بشكل طبيعي. يُعتبر المستوى 50 هو الخط الفاصل بين الاتجاه الصاعد والنازل."
        )

        # التطبيق والإشارات الذكية
        p3 = doc.add_paragraph()
        run3 = p3.add_run("\n⚙️ كيف تعمل المنظومة المالية الذكية؟\n")
        run3.font.name = 'Arial'
        run3.font.bold = True
        run3.font.size = Pt(14)
        run3.font.color.rgb = RGBColor(0, 51, 102)

        p3.add_run("تدمج المنصة مؤشر RSI مع المتوسط المتحرك (SMA 20) ومؤشر MACD لتوليد إشارات تداول ذكية:\n")
        doc.add_paragraph(
            "• 🟢 إشارة شراء (BUY): عندما يكون السعر فوق متوسط 20 يوماً مع RSI أقل من 45."
        )
        doc.add_paragraph(
            "• 🔴 إشارة بيع (SELL): عندما يكون السعر تحت متوسط 20 يوماً مع RSI أعلى من 55."
        )
        doc.add_paragraph(
            "• ⚪ إشارة انتظار (NEUTRAL): عندما تكون الحركة متذبذبة داخل النطاق الطبيعي."
        )

        # حفظ الملف في الذاكرة المؤقتة (Bytes)
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()